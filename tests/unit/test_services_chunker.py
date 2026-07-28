from __future__ import annotations

from mcp_cli.services.chunker import (
    chunk_by_sentences,
    chunk_by_tokens,
    extract_text,
    extract_text_from_pdf,
)


class TestChunkByTokens:
    def test_chunks_short_text(self):
        text = "hello world"
        chunks = chunk_by_tokens(text, chunk_size=10, overlap=2)
        assert len(chunks) == 1
        assert chunks[0]["text"] == "hello world"
        assert chunks[0]["word_count"] == 2

    def test_chunks_splits_at_token_boundary(self):
        words = "word " * 20
        chunks = chunk_by_tokens(words.strip(), chunk_size=5, overlap=1)
        assert len(chunks) >= 4
        assert all(c["word_count"] <= 5 for c in chunks)

    def test_chunks_with_overlap(self):
        words = "one two three four five six"
        chunks = chunk_by_tokens(words, chunk_size=3, overlap=1)
        assert len(chunks) >= 2
        if len(chunks) > 1:
            assert "three" in chunks[1]["text"]

    def test_empty_text_returns_empty_list(self):
        assert chunk_by_tokens("") == []
        assert chunk_by_tokens("   ") == []

    def test_single_chunk_for_small_text(self):
        text = "a b c d e"
        chunks = chunk_by_tokens(text, chunk_size=10, overlap=2)
        assert len(chunks) == 1
        assert chunks[0]["text"] == text

    def test_metadata_fields(self):
        chunks = chunk_by_tokens("a b c", chunk_size=2, overlap=0)
        assert "start_word" in chunks[0]
        assert "end_word" in chunks[0]
        assert "word_count" in chunks[0]


class TestChunkBySentences:
    def test_single_sentence(self):
        text = "Hello world."
        chunks = chunk_by_sentences(text, max_chars=500)
        assert len(chunks) == 1
        assert "Hello world." in chunks[0]["text"]

    def test_multiple_sentences_split(self):
        text = "A. B. C. D."
        chunks = chunk_by_sentences(text, max_chars=5)
        assert len(chunks) >= 2

    def test_empty_text(self):
        assert chunk_by_sentences("") == []
        assert chunk_by_sentences("   ") == []

    def test_metadata_in_chunks(self):
        text = "First sentence. Second sentence. Third one."
        chunks = chunk_by_sentences(text, max_chars=500)
        assert len(chunks) == 1
        assert "char_count" in chunks[0]
        assert "sentence_count" in chunks[0]

    def test_sentence_overlap(self):
        text = "Part one. Part two. Part three. Part four."
        chunks = chunk_by_sentences(text, max_chars=20, overlap_chars=10)
        assert len(chunks) >= 2
        if len(chunks) > 1:
            assert any("Part two" in c["text"] for c in chunks[1:])


class TestExtractText:
    def test_plain_text(self):
        content = b"hello world"
        assert extract_text(content, "test.txt") == "hello world"

    def test_plain_text_with_utf8(self):
        content = "héllo wörld 🎉".encode()
        result = extract_text(content, "notes.txt")
        assert "héllo" in result

    def test_unknown_extension_defaults_to_utf8(self):
        content = b"just some text"
        result = extract_text(content, "data.bin")
        assert result == "just some text"

    def test_pdf_extraction(self):
        content = _make_simple_pdf()
        result = extract_text(content, "doc.pdf")
        assert len(result) > 0

    def test_docx_extraction(self):
        content = _make_simple_docx()
        result = extract_text(content, "report.docx")
        assert len(result) > 0

    def test_empty_pdf_returns_empty(self):
        import io

        from pypdf import PdfWriter
        w = PdfWriter()
        buf = io.BytesIO()
        w.write(buf)
        result = extract_text_from_pdf(buf.getvalue())
        assert result == ""


def _make_simple_pdf() -> bytes:
    return b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (Hello PDF) Tj ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000362 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
433
%%EOF"""


def _make_simple_docx() -> bytes:
    import io

    from docx import Document
    doc = Document()
    doc.add_paragraph("Hello from docx.")
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
